import difflib
import io
import re
import logging
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from string import punctuation

from difference_between_files.acceptable import replacements, skips

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger('Documents')


class DiffData:
    def __init__(self, current_number, last_known_number, first_column, second_column, is_different):
        self.current_number = current_number
        self.last_known_number = last_known_number
        self.first_column = first_column
        self.second_column = second_column
        self.is_different = is_different


def list_from_string(document: str) -> list:
    lst = [i.replace('®', '').replace('', '') for i in document.split('\n\n')]
    return lst


def get_diff(list1: list, list2: list) -> list:
    last_known_number = ''
    for text1, text2 in zip(list1, list2):
        diff_items = difflib.ndiff([text1], [text2])
        logger.info('Find error')
        for diff in diff_items:
            diff = diff.strip('®')
            is_different = diff[0] in ('+', '-')
            break
        match_number = re.search(r"^(\.?,?\d{0,2}){0,4} ?", text1)
        current_number = match_number[0] if match_number and not match_number[0].isspace() else ""
        last_known_number = current_number if current_number else last_known_number
        diffs_data = DiffData(current_number=current_number, last_known_number=last_known_number, first_column=text1,
                              second_column=text2, is_different=is_different)

        yield diffs_data


def filter_diffs(diffs):
    duplicate = []
    for num, diff in enumerate(diffs):
        if not diff[1]:
            result = diff[0][:7]
            for n, dif in enumerate(diffs[num:]):
                number = n + num
                if result == dif[1][:7]:
                    duplicate.append((num, number, (diffs[number][0], diffs[num][1])))
                    break

    duplicate.reverse()

    for i in duplicate:
        diffs.pop(i[1])
        diffs[i[0]] = i[2]
    return diffs

def sequence_matcher(table,text1,text2,count_error,number):
    cells = table.add_row().cells
    cells[0].width = Inches(0.6)
    cells[1].width = Inches(3)
    cells[2].width = Inches(3)
    cells[0].text = number

    left_diff_count = 0
    right_diff_count = 0
    sequence = difflib.SequenceMatcher(a=text1.lower(), b=text2.lower(), autojunk=False)

    paragraph1 = cells[1].paragraphs[0]
    paragraph2 = cells[2].paragraphs[0]

    logger.info('Write document')
    for op, i1, i2, j1, j2 in sequence.get_opcodes():
        run1 = re.sub(u"[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+", '', text1[i1:i2])
        run2 = re.sub(u"[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+", '', text2[i1:i2])
        run1 = paragraph1.add_run(run1)
        run2 = paragraph2.add_run(run2)
        if op in ("delete", "insert"):
            if text1[i1:i2] not in skips:
                run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                left_diff_count += i2 - i1
            if text2[j1:j2] not in skips:
                run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                right_diff_count += j2 - j1
        if op == "replace":
            if (text1[i1:i2], text2[j1:j2]) not in replacements and (
                    text2[j1:j2],
                    text1[i1:i2],
            ) not in replacements:
                run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                left_diff_count += i2 - i1
                right_diff_count += j2 - j1

    if left_diff_count <= count_error and right_diff_count <= count_error and count_error>0:
        table._tbl.remove(table.rows[-1]._tr)


def add_paragraph(table,text1,text2,number):
    try:
        cells = table.add_row().cells
        cells[0].width = Inches(0.6)
        cells[1].width = Inches(3)
        cells[2].width = Inches(3)
        cells[0].text = number
        paragraph1 = cells[1].paragraphs[0]
        paragraph2 = cells[2].paragraphs[0]
        text1 = re.sub(u"[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+", '', text1)
        text2 = re.sub(u"[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+", '', text2)
        paragraph1.add_run(text1)
        paragraph2.add_run(text2)
        return
    except Exception as ex:
        print(ex)
        print()

def flagg_append(flag_text,text1,text2,number,diff):
    flag_text[1].append(text1)
    flag_text[2].append(text2)
    flag_text[3].append(diff.is_different)
    flag_text[0] = number.strip(' ✓ ')



def save_disagreement(file1: str, file2: str, count_error: int, flag: bool, file_name_docx: str, file_name_pdf: str) \
        -> io.BytesIO:
    logger.info('Create document')
    result = Document()
    result.add_heading("Протокол разногласий")
    table = result.add_table(rows=1, cols=3)
    table.style = "TableGrid"
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "№"
    heading_cells[0].width = Inches(0.6)
    heading_cells[1].text = f"Редакция заказчика\n{file_name_docx}"
    heading_cells[1].width = Inches(3)
    heading_cells[2].text = f"Редакция исполнителя\n{file_name_pdf}"
    heading_cells[2].width = Inches(3)

    list1, list2 = list_from_string(file1), list_from_string(file2)
    diffs = list(get_diff(list1, list2))
    number_flag = ''
    flag_text = [0, [], [],[]]
    for diff in diffs:
        if not diff.last_known_number and not diff.current_number:
            number = ''
        else:
            number = diff.last_known_number + ' ✓ ' if not diff.current_number else diff.current_number
        text1, text2 = [re.sub(r"(?:^(\.?,?\d{0,2}){0,4} ?|\.?,?$)", "", text).strip() for text in
                        [diff.first_column, diff.second_column]]
        if not diff.is_different and count_error == 0 and not flag:
            add_paragraph(table,text1,text2,number)
            continue
        if diff.is_different and count_error == 0 and not flag:
            sequence_matcher(table,text1,text2,count_error,number)
            continue

        if count_error == 0 and flag:
            if diff.current_number == '':
                flagg_append(flag_text,text1,text2,number,diff)

                continue
            else:
                if len(flag_text[1])>1 :
                    if  any([i for i in flag_text[3]]):
                        sequence_matcher(table, '\n'.join(flag_text[1]), '\n'.join(flag_text[2]), count_error, flag_text[0])
                        flag_text = [0, [], [],[]]
                        flagg_append(flag_text, text1, text2, number, diff)
                    else:
                        add_paragraph(table,'\n'.join(flag_text[1]),'\n'.join(flag_text[2]),flag_text[0])
                        flag_text = [0, [], [], []]
                        flagg_append(flag_text, text1, text2, number, diff)
                else:
                    if len(flag_text[1])==1 and not flag_text[3][0]:
                        add_paragraph(table, '\n'.join(flag_text[1]), '\n'.join(flag_text[2]), flag_text[0])
                        flag_text = [0, [], [],[]]
                        flagg_append(flag_text, text1, text2, number, diff)
                    else:
                        if len(flag_text[1]) == 1 and flag_text[3][0]:
                            sequence_matcher(table, '\n'.join(flag_text[1]), '\n'.join(flag_text[2]), count_error,
                                             flag_text[0])
                            flag_text = [0, [], [], []]
                            flagg_append(flag_text, text1, text2, number, diff)

                        else:
                            flagg_append(flag_text, text1, text2, number, diff)

                continue
        if count_error>0 and not flag:
            sequence_matcher(table,text1,text2,count_error,number)
            continue

        if count_error>0 and flag:
                if diff.current_number == '':
                    flagg_append(flag_text, text1, text2, number, diff)
                    continue
                if len(flag_text[1])>1 :
                    if  all([not i for i in flag_text[3]]):
                        flag_text = [0, [], [], []]
                        continue
                    else:
                        sequence_matcher(table,'\n'.join(flag_text[1]),'\n'.join(flag_text[2]),count_error,flag_text[0])
                        flag_text = [0, [], [],[]]
                        flagg_append(flag_text, text1, text2, number, diff)
                else:
                    if len(flag_text[1])==1 and flag_text[3][0]:
                        sequence_matcher(table,flag_text[1][0],flag_text[2][0],count_error,flag_text[0])
                        flag_text = [0, [], [], []]
                        flagg_append(flag_text, text1, text2, number, diff)
                    else:
                        flag_text = [0, [], [], []]
                        flagg_append(flag_text, text1, text2, number, diff)


    else:
        if len(flag_text[1])>0:
            sequence_matcher(table, '\n'.join(flag_text[1]), '\n'.join(flag_text[2]), count_error, flag_text[0])




    file_stream = io.BytesIO()
    result.save(file_stream)
    file_stream.seek(0)

    return file_stream
