import re
import subprocess
import fuzzywuzzy
import pdfplumber
from __init__ import *
from typing import List
from docx import Document
from collections import deque
from fuzzywuzzy import process, fuzz
from docx2python import docx2python

logging.getLogger("pdfminer").setLevel(logging.WARNING)


class Docx(object):
    def __init__(self, absolute_path_filename: str):
        self.absolute_path_filename: str = absolute_path_filename

    @staticmethod
    def clean_special_chars(lst: List[str]) -> List[str]:
        lst = [s.replace(u"\u202F", " ") for s in lst]
        lst = [re.sub(' +', ' ', s) for s in lst]
        lst = [re.sub('--\n', '', s) for s in lst]
        lst = [re.sub('--\t', '', s) for s in lst]
        lst = [re.sub('\t', '', s) for s in lst]
        lst = [s for s in lst if s.strip()]
        return lst

    @staticmethod
    def get_paragraph_starts(docx_text: List[str], pdf_text: List[str]) -> List[int]:
        # index_paragraph = 0
        # p_starts = set()
        p_starts = deque([0, 0, 0])
        for dl in docx_text:
            if len(dl) < 3:
                continue
            i = min(70, len(dl))
            prefix = dl[:i]
            # pls = process.extract(prefix, pdf_text[index_paragraph:], limit=1)
            pls = fuzzywuzzy.process.extract(prefix,
                                             pdf_text[min(list(p_starts)[-3:]):],
                                             scorer=fuzzywuzzy.fuzz.ratio,
                                             limit=3)
            pl = pls[0][0]
            new_index = pdf_text.index(pl)
            p_starts.append(new_index)
        return sorted(set(p_starts))
        #     index_paragraph = pdf_text.index(pl)
        #     p_starts.add(index_paragraph)
        # return sorted(p_starts)

    def save_txt(self, paragraphs):
        path = os.path.join(os.path.dirname(self.absolute_path_filename), "paragraphs.txt")
        with open(path, "w") as f:
            f.writelines(paragraphs)

    def format_paragraphs(self, docx_text: List[str], pdf_text: List[str]) -> str:
        docx_text = self.clean_special_chars(docx_text)
        pdf_text = self.clean_special_chars(pdf_text)
        p_starts = self.get_paragraph_starts(docx_text, pdf_text)
        paragraphs = []
        current_paragraph = ''
        for i, pl in enumerate(pdf_text):
            if i in p_starts:
                if current_paragraph:
                    paragraphs.append(current_paragraph)
                current_paragraph = pl
            else:
                current_paragraph = current_paragraph.replace("\n", " ")
                current_paragraph += pl
        paragraphs.append(current_paragraph)
        self.save_txt(paragraphs)
        return "".join(paragraphs)

    def refactor_page_header(self, flag) -> None:
        if flag:
            doc = Document(self.absolute_path_filename)
            doc.sections[0].header.paragraphs[0].text = ''
            doc.sections[0].footer.paragraphs[0].text = ''
            doc.save(self.absolute_path_filename)

    def convert_to_docx(self) -> None:
        subprocess.check_output(['libreoffice', '--convert-to', 'docx', self.absolute_path_filename, '--outdir',
                                 os.path.dirname(self.absolute_path_filename)])
        self.absolute_path_filename += 'x'

    def get_text(self, mime_type) -> str:
        if "Composite Document File V2 Document" in mime_type:
            self.convert_to_docx()
        # self.refactor_page_header(True)
        docx_text = docx2python(self.absolute_path_filename)
        subprocess.check_output(['libreoffice', '--convert-to', 'pdf', self.absolute_path_filename, '--outdir',
                                 os.path.dirname(self.absolute_path_filename)])
        texts = pdfplumber.open(self.absolute_path_filename.replace(".docx", ".pdf"))
        list_pdf_text = []
        for text in texts.pages:
            list_pdf_text.extend(line.strip() + '\n' for line in text.extract_text().split('\n'))
        list_docx_text = [line.strip() + '\n' for line in docx_text.text.split('\n')]
        with open(f"{os.path.dirname(self.absolute_path_filename)}/list_pdf_text.txt", "w") as f:
            f.writelines(list_pdf_text)
        with open(f"{os.path.dirname(self.absolute_path_filename)}/list_docx_text.txt", "w") as f:
            f.writelines(list_docx_text)
        return self.format_paragraphs(list_docx_text, list_pdf_text)
